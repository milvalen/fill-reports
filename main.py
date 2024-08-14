import ezodf
from docx import Document
import os


def read_ods(file_path):
    doc = ezodf.opendoc(file_path)
    sheet = doc.sheets[0]

    placeholders = {}
    for row in sheet.rows():
        if len(row) >= 2:
            key = str(row[0].value).strip()
            value = str(row[1].value).strip()
            placeholders[key] = value
    
    return placeholders


def replace_placeholders(doc, placeholders):
    ds_number = int(float(placeholders["ДС"]))
    contract_number = placeholders.get("Номер договора", "без номера")
    ds_replacements = generate_ds_strings(contract_number, ds_number)
    insert_ds_replacements(doc, ds_replacements)
    del placeholders["ДС"]

    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if f"<<<{placeholder}>>>" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"<<<{placeholder}>>>", value)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in placeholders.items():
                    if f"<<<{placeholder}>>>" in cell.text:
                        cell.text = cell.text.replace(f"<<<{placeholder}>>>", value)
                    if "<<<ДС>>>" in cell.text:
                        cell.text = cell.text.replace("<<<ДС>>>", "")


def generate_ds_strings(contract_number, ds_number):
    return [f"ДС № {contract_number}/{i}" for i in range(1, ds_number + 1)]


def insert_ds_replacements(doc, ds_replacements):
    for paragraph in doc.paragraphs:
        if "<<<ДС>>>" in paragraph.text:
            for ds_string in ds_replacements:
                paragraph.add_run(ds_string).add_break()

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "<<<ДС>>>" in cell.text:
                    for ds_string in ds_replacements:
                        cell.add_paragraph(ds_string)


def fill_template(ods_file, docx_template, output_folder):
    placeholders = read_ods(ods_file)
    
    doc = Document(docx_template)
    
    replace_placeholders(doc, placeholders)
    
    contract_number = placeholders.get("Номер договора", "без номера")
    output_filename = f"ЭЗ {contract_number}.docx"
    output_path = os.path.join(output_folder, output_filename)
    
    doc.save(output_path)
    print(f"Документ сохранен как: {output_path}")

if __name__ == '__main__':
    ods_file = "blank.ods"
    docx_template = "Balvanka.docx"
    output_folder = "filled"

    os.makedirs(output_folder, exist_ok=True)

    fill_template(ods_file, docx_template, output_folder)
